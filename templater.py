import re
import sys
import os
import openpyxl as xl
import openpyxl.utils.exceptions as ex
from docx import Document
from docx.shared import Inches




def docx_replace_regex(doc_obj, regex, replace):

    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex, replace)



def fill_docs(data_filename, template_filename,path):
    MARKERS = []
    f = open("static/doc_example/markers.txt", "r")
    for x in f:
        MARKERS.append(re.compile(r''+str(x).strip()))
    try:
        data = xl.load_workbook(data_filename)
    except ex.InvalidFileException:
        print("Wrong file")
        sys.exit()
    except TypeError:
        sys.exit()
    ws = data.worksheets[0]
    doc = Document(template_filename)
    if not os.path.exists(path):
        os.mkdir(path)
    for r in ws.rows:
        r = list(r)
        for i, v in enumerate(MARKERS):
            docx_replace_regex(doc, v, r''+str(r[i].value))
        doc.tables[1].rows[0].cells[1].paragraphs[0].add_run().add_picture(
        'static/doc_example/' + str(r[4].value) + '.png', width=Inches(1.6), height=Inches(0.8)) 
        doc.save('docs/' + r[1].value + '.docx')
        doc = Document(template_filename)
